from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent.parent
REPORT_MD_PATH = ROOT / "report" / "report.md"
ASSETS_DIR = ROOT / "report" / "assets"
DOCX_PATH = ROOT / "report" / "股指期货套利策略研究报告.docx"
DOCX_FALLBACK_PATH = ROOT / "report" / "股指期货套利策略研究报告_v2.docx"

BODY_COLOR = RGBColor(0x22, 0x22, 0x22)
MUTED_COLOR = RGBColor(0x55, 0x55, 0x55)
ACCENT_COLOR = RGBColor(0x00, 0x00, 0x00)
HEADER_FILL = "EDEDED"
ROW_FILL = "F7F7F7"
INFO_LEFT_FILL = "EEEEEE"
INFO_RIGHT_FILL = "FAFAFA"
BODY_LINE_SPACING_PT = 20
MAX_IMAGE_WIDTH_CM = 15.2
MAX_IMAGE_HEIGHT_CM = 18.8

FIGURE_CAPTIONS = {
    "fig_roll_yield.png": (
        "主力合约持有段收益分布",
        "蓝色为期货持有收益，红色为扣减同期指数收益后的超额部分。正超额占比在 50% 附近，单次换仓不确定性大，但长期均值仍为正。",
    ),
    "fig_nav.png": (
        "策略净值曲线",
        "比较吃贴水策略与现货指数的净值走势。Always 策略在 IC 和 IM 上都明显高于指数基准，但不是平滑套利曲线。",
    ),
    "fig_annual.png": (
        "年度收益拆分",
        "按自然年比较 Always、B<-1.5% 与指数基准的收益表现，用于观察收益是否集中在少数年份。",
    ),
    "fig_drawdown.png": (
        "回撤路径",
        "展示策略与指数基准的回撤深度和修复过程。回撤曲线用于判断策略是否适合实际资金承受。",
    ),
    "fig_heatmap.png": (
        "参数敏感性热力图",
        "展示基差阈值与展期窗口对 Sharpe 和 MaxDD 的影响。展期窗口影响不大，基差阈值影响更明显，但为样本内观察。",
    ),
    "fig_calendar_leverage.png": (
        "跨期套利杠杆压力测试",
        "在同一条跨期收益序列上测试 1x-5x 杠杆，展示收益、回撤和保证金占用的变化。",
    ),
    "fig_bootstrap.png": (
        "Bootstrap 显著性检验",
        "Block Bootstrap（block=10天）产生的 SR 差异分布。IC 和 IM 的观测 SR 差异均为正，95% CI 不覆盖 0。",
    ),
    "fig_attribution.png": (
        "收益归因分解（无截距）",
        "将策略年化收益按无截距模型拆分为 beta 和残差，比较方向暴露与基差/展期补偿的相对贡献。",
    ),
    "fig_cost_stress.png": (
        "成本压力测试",
        "将交易成本从 1 倍逐级放大到 10 倍，观察策略年化收益的变化。策略对成本不敏感。",
    ),
    "fig_correlation.png": (
        "IC-IM 日收益相关性",
        "IC 与 IM 吃贴水策略日收益的散点图和回归线。相关性约 0.95，品种分散化价值有限。",
    ),
}

TABLE_CAPTIONS = [
    ("研究口径与核心设定", "概括报告的研究标的、数据区间、交易假设和评价重点。"),
    ("数据质量检查", "对比 IC 与 IM 的原始样本量、主力合约样本量、指数匹配率和合约切换统计。"),
    ("回测假设与交易口径", "统一所有策略的信号时点、成交时点、合约选择、展期规则和成本设定。"),
    ("基差因子前瞻 IC 分析", "ADF 均值回复检验、半衰期、基差对期货和指数前瞻 1/5/20 日收益的预测 IC。"),
    ("主力合约持有段收益统计", "IC 与 IM 各持有段的期货收益和超额收益分布。"),
    ("吃贴水策略主要绩效", "比较 Always、B<-1.5% 和指数基准在 IC 与 IM 上的收益、风险与交易统计。"),
    ("跨期套利回测结果", "代表性参数下 IC 和 IM 的跨期套利绩效。"),
    ("跨期套利杠杆压力测试", "在同一条跨期策略收益序列上测试 1x-5x 杠杆，展示收益、回撤和保证金占用的变化。"),
    ("吃贴水+跨期套利组合绩效", "展示 Carry Only、Spread Only、90/10、80/20、70/30 五种权重的组合绩效及相对 Carry Only 的改善。"),
    ("Block Bootstrap 显著性检验", "循环块 Bootstrap 重采样下策略相对指数 Sharpe 差异的置信区间和 p 值。"),
    ("无截距收益归因", "将策略日收益按无截距模型拆解为指数驱动和残差两部分。"),
    ("图表清单", "报告中全部图表的编号、文件名和内容说明。"),
]


def set_run_font(
    run,
    size_pt: float = 12,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    latin_font: str = "Times New Roman",
    east_asia_font: str = "宋体",
):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = latin_font
    run.font.color.rgb = color or BODY_COLOR

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin_font)
    rfonts.set(qn("w:hAnsi"), latin_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def clean_inline(text: str) -> str:
    text = text.replace("<br>", " ").replace("<br/>", " ").replace("<br />", " ")
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def add_inline_runs(
    paragraph,
    text: str,
    *,
    size_pt: float = 12,
    base_bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    east_asia_font: str = "宋体",
):
    text = clean_inline(text)
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(
                run, size_pt=size_pt, bold=base_bold, italic=italic,
                color=color, east_asia_font=east_asia_font,
            )

        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2].replace("`", ""))
            set_run_font(
                run, size_pt=size_pt, bold=True, italic=italic,
                color=color, east_asia_font=east_asia_font,
            )
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run, size_pt=size_pt - 0.5, color=MUTED_COLOR,
                latin_font="Consolas", east_asia_font="宋体",
            )
        pos = match.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(
            run, size_pt=size_pt, bold=base_bold, italic=italic,
            color=color, east_asia_font=east_asia_font,
        )


def set_paragraph_bottom_border(paragraph, color: str = "808080", size: str = "8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, left: int = 80, right: int = 80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_compact_spacing(paragraph, *, before: float = 0, after: float = 0):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def add_word_field(paragraph, instruction: str, placeholder: str = ""):
    """Insert a Word field, such as TOC or PAGE, using raw OOXML."""
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)
    run._r.append(end)
    set_run_font(run, size_pt=10.5)
    return run


def enable_update_fields_on_open(doc: Document):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)

    head = paragraph.add_run("第 ")
    set_run_font(head, size_pt=9.5, color=MUTED_COLOR)

    add_word_field(paragraph, " PAGE ", "1")

    tail = paragraph.add_run(" 页")
    set_run_font(tail, size_pt=9.5, color=MUTED_COLOR)


def add_header(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("股指期货套利策略研究")
    set_run_font(run, size_pt=9.5, color=MUTED_COLOR)


def configure_document(doc: Document):
    enable_update_fields_on_open(doc)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = BODY_COLOR
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    heading_specs = {
        "Heading 1": (14.5, "黑体", 12, 6),
        "Heading 2": (12.5, "黑体", 8, 4),
        "Heading 3": (11.5, "黑体", 6, 3),
    }
    for style_name, (size_pt, east_font, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size_pt)
        style.font.bold = True
        style.font.color.rgb = ACCENT_COLOR
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_font)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.line_spacing = 1

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.55)
    section.right_margin = Cm(2.55)

    add_header(section.header.paragraphs[0])
    add_page_number(section.footer.paragraphs[0])


def add_horizontal_rule(doc: Document, color: str = "000000", size: str = "10"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_bottom_border(p, color=color, size=size)


def add_cover_page(doc: Document):
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("股指期货套利策略研究报告")
    set_run_font(run, size_pt=24, bold=True, color=ACCENT_COLOR, east_asia_font="黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("吃贴水与跨期套利的策略评估与风险拆解")
    set_run_font(run, size_pt=15, color=MUTED_COLOR, east_asia_font="黑体")

    add_horizontal_rule(doc)

    items = [
        ("研究标的", "IC（中证500股指期货）与 IM（中证1000股指期货）"),
        ("数据区间", "IC：2015-04-16 至 2026-03-02；IM：2022-07-22 至 2026-03-02"),
        ("策略口径", "吃贴水多头 + 基差阈值择时 + 跨期价差套利"),
        ("评估维度", "收益、风险、统计显著性、交易可执行性、可复现性"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for left, right in items:
        row = table.add_row().cells
        set_cell_shading(row[0], INFO_LEFT_FILL)
        set_cell_shading(row[1], INFO_RIGHT_FILL)
        for cell in row:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = row[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(left)
        set_run_font(r0, size_pt=11, bold=True, east_asia_font="黑体")
        p1 = row[1].paragraphs[0]
        add_inline_runs(p1, right, size_pt=11)

    for _ in range(7):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("金融工程策略研究")
    set_run_font(run, size_pt=12, color=MUTED_COLOR)

    doc.add_page_break()


def add_navigation_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("目录")
    set_run_font(run, size_pt=18, bold=True, east_asia_font="黑体")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_word_field(p, ' TOC \\o "1-2" \\h \\z \\u ', "打开文档后目录将自动更新")

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int):
    if level == 2 and re.match(r"^\d+\.\s", text) and not text.startswith("1. "):
        doc.add_section(WD_SECTION_START.NEW_PAGE)

    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(16)
        add_inline_runs(p, text, size_pt=20, base_bold=True, east_asia_font="黑体")
        return

    if level == 2:
        word_level = 1
    elif level == 3 and re.match(r"^\d+\.\d+", text):
        word_level = 2
    else:
        word_level = 3
    p = doc.add_paragraph(style=f"Heading {word_level}")
    p.paragraph_format.space_before = Pt(12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(6 if level in (2, 3) else 3)
    if level == 2:
        add_inline_runs(p, text, size_pt=14.5, base_bold=True, east_asia_font="黑体")
        set_paragraph_bottom_border(p)
    elif level == 3:
        add_inline_runs(p, text, size_pt=12.5, base_bold=True, east_asia_font="黑体")
    else:
        add_inline_runs(p, text, size_pt=11.5, base_bold=True, color=MUTED_COLOR, east_asia_font="黑体")


def add_text_paragraph(doc: Document, text: str, *, center: bool = False, size_pt: float = 12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    p.paragraph_format.space_after = Pt(6)
    if not center:
        p.paragraph_format.first_line_indent = Cm(0.74)
    add_inline_runs(p, text, size_pt=size_pt)


def add_lead_line(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text, size_pt=11.5, italic=True, color=MUTED_COLOR)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    p.paragraph_format.space_after = Pt(3)
    add_inline_runs(p, text, size_pt=12)


def add_numbered_item(doc: Document, number: str, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.44)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    p.paragraph_format.space_after = Pt(3)
    add_inline_runs(p, f"{number}. {text}", size_pt=12)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_compact_spacing(p, after=6)
    add_inline_runs(p, text, size_pt=10.5, base_bold=True, color=MUTED_COLOR)


def add_figure_caption(doc: Document, figure_no: int, title: str):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.keep_together = True
    title_p.paragraph_format.keep_with_next = True
    set_compact_spacing(title_p, after=4)
    add_inline_runs(title_p, f"图 {figure_no} {title}", size_pt=10.5, base_bold=True, color=ACCENT_COLOR)


def add_table_caption(doc: Document, table_no: int):
    if table_no <= len(TABLE_CAPTIONS):
        title, note = TABLE_CAPTIONS[table_no - 1]
    else:
        title, note = "补充表格", "汇总报告正文中的补充数据。"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    set_compact_spacing(p, before=4, after=4)
    add_inline_runs(p, f"表 {table_no} {title}：{note}", size_pt=10.5, base_bold=True, color=MUTED_COLOR)


def add_source_line(doc: Document, text: str = "资料来源：IC/IM 期货与指数数据，本文计算"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_compact_spacing(p, before=0, after=6)
    add_inline_runs(p, text, size_pt=8.5, color=MUTED_COLOR)


def resolve_image(rel_path: str) -> Path:
    raw = rel_path.strip().replace("\\", "/")
    path = REPORT_MD_PATH.parent / raw
    if path.exists():
        return path
    path = ASSETS_DIR / Path(raw).name
    if path.exists():
        return path
    raise FileNotFoundError(f"Image referenced in report.md was not found: {rel_path}")


def image_display_size_cm(img_path: Path) -> tuple[float, float | None]:
    with Image.open(img_path) as img:
        width_px, height_px = img.size
    height_at_max_width = MAX_IMAGE_WIDTH_CM * height_px / width_px
    if height_at_max_width <= MAX_IMAGE_HEIGHT_CM:
        return MAX_IMAGE_WIDTH_CM, None
    width_at_max_height = MAX_IMAGE_HEIGHT_CM * width_px / height_px
    return width_at_max_height, MAX_IMAGE_HEIGHT_CM


def add_image(doc: Document, img_path: Path, figure_no: int, alt_text: str = ""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    set_compact_spacing(p, before=4, after=2)
    run = p.add_run()
    width_cm, height_cm = image_display_size_cm(img_path)
    if height_cm is None:
        run.add_picture(str(img_path), width=Cm(width_cm))
    else:
        run.add_picture(str(img_path), height=Cm(height_cm))

    title, _note = FIGURE_CAPTIONS.get(
        img_path.name,
        (alt_text or img_path.stem, "展示报告正文对应分析结果。"),
    )
    add_figure_caption(doc, figure_no, title)
    add_source_line(doc)


def is_table_block(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    current = lines[idx].strip()
    nxt = lines[idx + 1].strip()
    return current.startswith("|") and nxt.startswith("|") and set(nxt.replace("|", "").strip()) <= {"-", ":", " "}


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_markdown_table(lines: list[str], idx: int):
    header = split_table_row(lines[idx])
    rows = []
    i = idx + 2
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        row = split_table_row(line)
        if len(row) < len(header):
            row += [""] * (len(header) - len(row))
        rows.append(row[: len(header)])
        i += 1
    return header, rows, i


def sanitize_table_text(text: str) -> str:
    text = clean_inline(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    return text.replace("`", "")


def add_table(doc: Document, header: list[str], rows: list[list[str]], table_no: int):
    add_table_caption(doc, table_no)
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    cols = len(header)
    font_size = 10.5 if cols <= 3 else 9.5 if cols <= 6 else 8.5
    is_info_table = [sanitize_table_text(x) for x in header] == ["项目", "内容"] or \
                    [sanitize_table_text(x) for x in header] == ["项目", "口径"]

    for j, text in enumerate(header):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell, top=70, bottom=70, left=60, right=60)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        set_compact_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(
            p,
            sanitize_table_text(text),
            size_pt=font_size,
            base_bold=True,
            east_asia_font="黑体",
        )

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            cell = cells[j]
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if is_info_table:
                set_cell_shading(cell, INFO_LEFT_FILL if j == 0 else INFO_RIGHT_FILL)
            elif row_idx % 2 == 1:
                set_cell_shading(cell, ROW_FILL)
            p = cell.paragraphs[0]
            set_compact_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, sanitize_table_text(text), size_pt=font_size)

    spacer = doc.add_paragraph()
    set_compact_spacing(spacer, after=2)
    add_source_line(doc)


def build_docx() -> Path:
    if not REPORT_MD_PATH.exists():
        raise FileNotFoundError(f"Cannot find {REPORT_MD_PATH}")

    lines = REPORT_MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_cover_page(doc)
    add_navigation_page(doc)

    paragraph_buffer: list[str] = []
    figure_no = 0
    table_no = 0

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(x.strip() for x in paragraph_buffer if x.strip())
            if text:
                add_text_paragraph(doc, text)
            paragraph_buffer = []

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped in {"---", "***", "___"}:
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue

        if stripped.startswith("#### "):
            flush_paragraph()
            add_heading(doc, stripped[5:].strip(), 4)
            i += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            add_lead_line(doc, stripped[2:].strip())
            i += 1
            continue

        if is_table_block(lines, i):
            flush_paragraph()
            header, rows, next_idx = parse_markdown_table(lines, i)
            table_no += 1
            add_table(doc, header, rows, table_no)
            i = next_idx
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            flush_paragraph()
            alt_text = image_match.group(1).strip()
            img_path = resolve_image(image_match.group(2))
            figure_no += 1
            add_image(doc, img_path, figure_no, alt_text=alt_text)
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            add_bullet(doc, stripped[2:].strip())
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if number_match:
            flush_paragraph()
            add_numbered_item(doc, number_match.group(1), number_match.group(2).strip())
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()

    try:
        doc.save(DOCX_PATH)
        return DOCX_PATH
    except PermissionError:
        doc.save(DOCX_FALLBACK_PATH)
        return DOCX_FALLBACK_PATH


if __name__ == "__main__":
    output = build_docx()
    print(f"Generated {output.relative_to(ROOT)}")
